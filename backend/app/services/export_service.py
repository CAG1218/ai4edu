"""
AI4Edu 导出服务
笔记导出PDF/Word、诊断报告导出、数据批量导出
"""
import io
import json
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from sqlalchemy import and_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import NotFoundException
from app.models.note import Note
from app.models.diagnosis import Diagnosis

logger = logging.getLogger(__name__)


class ExportService:
    """导出服务"""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def export_note_pdf(
        self,
        note_id: int,
        tenant_id: int,
    ) -> bytes:
        """
        导出笔记为PDF

        Args:
            note_id: 笔记ID
            tenant_id: 租户ID

        Returns:
            PDF文件二进制数据
        """
        note_data = await self._get_note_data(note_id, tenant_id)
        if not note_data:
            raise NotFoundException(message="笔记不存在")

        # 生成PDF（使用reportlab或简单HTML→PDF）
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()

            # 自定义中文样式
            title_style = ParagraphStyle(
                "ChineseTitle",
                parent=styles["Title"],
                fontSize=18,
                spaceAfter=12,
            )
            body_style = ParagraphStyle(
                "ChineseBody",
                parent=styles["Normal"],
                fontSize=12,
                leading=18,
                spaceAfter=6,
            )

            elements = []

            # 标题
            elements.append(Paragraph(note_data["title"], title_style))
            elements.append(Spacer(1, 12))

            # 内容
            content = note_data.get("content", "") or note_data.get("content_plain", "")
            if content:
                # 简单的Markdown→HTML→Paragraph转换
                paragraphs = content.split("\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        # 转义HTML特殊字符
                        safe_text = para_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        elements.append(Paragraph(safe_text, body_style))

            doc.build(elements)
            return buffer.getvalue()

        except ImportError:
            logger.warning("reportlab未安装，使用纯文本回退")
            content = f"# {note_data['title']}\n\n{note_data.get('content', '') or note_data.get('content_plain', '')}"
            return content.encode("utf-8")

    async def export_note_docx(
        self,
        note_id: int,
        tenant_id: int,
    ) -> bytes:
        """
        导出笔记为Word文档

        Args:
            note_id: 笔记ID
            tenant_id: 租户ID

        Returns:
            Word文件二进制数据
        """
        note_data = await self._get_note_data(note_id, tenant_id)
        if not note_data:
            raise NotFoundException(message="笔记不存在")

        try:
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()

            # 标题
            heading = doc.add_heading(note_data["title"], level=1)

            # 内容
            content = note_data.get("content", "") or note_data.get("content_plain", "")
            if content:
                paragraphs = content.split("\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)

            # 元数据
            doc.add_paragraph("")
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"创建时间: {note_data.get('created_at', '')}").italic = True

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            logger.warning("python-docx未安装，使用纯文本回退")
            content = f"# {note_data['title']}\n\n{note_data.get('content', '') or note_data.get('content_plain', '')}"
            return content.encode("utf-8")

    async def export_diagnosis_report(
        self,
        diagnosis_id: int,
        tenant_id: int,
    ) -> bytes:
        """
        导出诊断报告

        Args:
            diagnosis_id: 诊断ID
            tenant_id: 租户ID

        Returns:
            报告文件二进制数据（Markdown格式）
        """
        stmt = select(Diagnosis).where(
            and_(
                Diagnosis.id == diagnosis_id,
                Diagnosis.tenant_id == tenant_id,
            )
        )
        result = await self.db.execute(stmt)
        diagnosis = result.scalars().first()

        if not diagnosis:
            raise NotFoundException(message="诊断不存在")

        # 生成Markdown报告
        weaknesses = []
        if diagnosis.weaknesses:
            try:
                weaknesses = json.loads(diagnosis.weaknesses)
            except (json.JSONDecodeError, TypeError):
                pass

        strengths = []
        if diagnosis.strengths:
            try:
                strengths = json.loads(diagnosis.strengths)
            except (json.JSONDecodeError, TypeError):
                pass

        recommendations = []
        if diagnosis.recommendations:
            try:
                recommendations = json.loads(diagnosis.recommendations)
            except (json.JSONDecodeError, TypeError):
                pass

        report_lines = [
            f"# {diagnosis.title}",
            "",
            f"**诊断类型**: {diagnosis.diagnosis_type}",
            f"**状态**: {diagnosis.status}",
            f"**得分**: {diagnosis.score}分" if diagnosis.score else "**得分**: 未评分",
            f"**正确率**: {diagnosis.correct_count}/{diagnosis.total_questions}",
            f"**用时**: {diagnosis.duration_seconds}秒" if diagnosis.duration_seconds else "",
            "",
            "## 知识强项",
            "",
        ]

        for s in strengths:
            report_lines.append(f"- ✅ {s}")

        report_lines.extend([
            "",
            "## 知识弱点",
            "",
        ])

        for w in weaknesses:
            report_lines.append(f"- ⚠️ {w}")

        report_lines.extend([
            "",
            "## 学习建议",
            "",
        ])

        for r in recommendations:
            report_lines.append(f"- 📚 {r}")

        report_lines.extend([
            "",
            "---",
            f"*报告生成时间: {datetime.utcnow().isoformat()}*",
        ])

        return "\n".join(report_lines).encode("utf-8")

    async def export_user_data_batch(
        self,
        tenant_id: int,
        user_id: int,
        data_types: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        批量导出用户数据

        Args:
            tenant_id: 租户ID
            user_id: 用户ID
            data_types: 要导出的数据类型列表

        Returns:
            用户数据包
        """
        if data_types is None:
            data_types = ["notes", "diagnoses"]

        export_data: Dict[str, Any] = {
            "user_id": user_id,
            "tenant_id": tenant_id,
            "exported_at": datetime.utcnow().isoformat(),
            "data": {},
        }

        # 导出笔记
        if "notes" in data_types:
            notes_stmt = select(Note).where(
                and_(
                    Note.owner_id == user_id,
                    Note.tenant_id == tenant_id,
                    Note.is_deleted == False,
                )
            )
            notes_result = await self.db.execute(notes_stmt)
            notes = notes_result.scalars().all()
            export_data["data"]["notes"] = [
                {
                    "id": n.id,
                    "title": n.title,
                    "content": n.content,
                    "note_type": n.note_type,
                    "word_count": n.word_count,
                    "created_at": n.created_at.isoformat() if n.created_at else None,
                }
                for n in notes
            ]

        # 导出诊断
        if "diagnoses" in data_types:
            diag_stmt = select(Diagnosis).where(
                and_(
                    Diagnosis.user_id == user_id,
                    Diagnosis.tenant_id == tenant_id,
                )
            )
            diag_result = await self.db.execute(diag_stmt)
            diagnoses = diag_result.scalars().all()
            export_data["data"]["diagnoses"] = [
                {
                    "id": d.id,
                    "title": d.title,
                    "diagnosis_type": d.diagnosis_type,
                    "score": d.score,
                    "status": d.status,
                    "created_at": d.created_at.isoformat() if d.created_at else None,
                }
                for d in diagnoses
            ]

        return export_data

    async def _get_note_data(
        self,
        note_id: int,
        tenant_id: int,
    ) -> Optional[Dict[str, Any]]:
        """获取笔记数据"""
        stmt = select(Note).where(
            and_(
                Note.id == note_id,
                Note.tenant_id == tenant_id,
                Note.is_deleted == False,
            )
        )
        result = await self.db.execute(stmt)
        note = result.scalars().first()

        if not note:
            return None

        tags = []
        if note.tags:
            try:
                tags = json.loads(note.tags)
            except (json.JSONDecodeError, TypeError):
                pass

        return {
            "id": note.id,
            "title": note.title,
            "content": note.content,
            "content_plain": note.content_plain,
            "note_type": note.note_type,
            "tags": tags,
            "word_count": note.word_count,
            "version": note.version,
            "created_at": note.created_at.isoformat() if note.created_at else None,
            "updated_at": note.updated_at.isoformat() if note.updated_at else None,
        }
